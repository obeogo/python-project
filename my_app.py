"""
class Person:
    def __init__(self, name, age):
        self.name = name
        self.age = age

    def walk(self):
        print(self.name + ' is walking...')

    def speak(self):
        print('Hello! my name is ' + self.name + ' and I am ' + str(self.age) + ' years old')

john = Person('John', 22)
john.speak()
john.walk()

print('----------------------------------')

mariam = Person('Mariam', 18)
mariam.speak()
mariam.walk()
"""


from docx import Document
from docx.shared import Inches
import pyttsx3


def speak(text):
    pyttsx3.speak(text)

document = Document()

#PROFILE PICTURE
document.add_picture(
    'me.jpg', 
    width=Inches(2.0))

#NAME PHONE EMAIL DETAILS
name = input('what is your name? ')
speak('Hello ' + name + 'how are you today?')

speak('what is your phone number? ')
phone_number = input('what is your phone number? ')
email = input('what is your email? ')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email)

#ABOUT ME
document.add_heading('About me')
about_me = input('Tell me about yourself? ')
document.add_paragraph(about_me)

#WORK EXPERIENCE
document.add_heading('Work Experience ')
p = document.add_paragraph()

company = input('Enter company ')
from_date = input('Fram Date ')
to_date = input('To Date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input(
    'Describe your experience at ' + company)
p.add_run(experience_details)

# MORE EXPERIENCES
while True:
    has_more_experiences = input(
        'Do you have more experiences? Yes or No ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter company ')
        from_date = input('Fram Date ')
        to_date = input('To Date ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input(
            'Describe your experience at ' + company + ' ')
        p.add_run(experience_details)
    else:
        break

# SKILLS
document.add_heading('Skills')
skill = input('Enter skill')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have more skills? Yes or No ')
    if has_more_skills.lower() == 'yes':
        skill = input('Enter skill')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

#FOOTER
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated by Ousmane Beogo: Python Developer at Foundeas Technology for Google."


document.save('cv.docx')
